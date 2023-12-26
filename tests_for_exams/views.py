import random

from django.http import HttpResponse
from docx import Document
from django.contrib import messages
from django.shortcuts import redirect
from django.views import View
from django.views.generic import ListView, TemplateView, DetailView
from random import sample

from tests_for_exams.models import Question, AnswerOption, CorrectAnswer, Test, UserAnswer, Subject


# Create your views here.
class Home(TemplateView):
    template_name = 'home.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['tests'] = Test.objects.all()
        return context


class QuestionListView(ListView):
    model = Question
    template_name = 'questions_list.html'
    context_object_name = 'questions'
    paginate_by = 10



    def post(self, request):
        for key, value in request.POST.items():
            if key.startswith('correct_answer_'):
                question_id = key.split('_')[-1]
                answer_id = value
                question = Question.objects.get(id=question_id)
                answer = AnswerOption.objects.get(id=answer_id)

                CorrectAnswer.objects.update_or_create(question=question, defaults={'answer': answer})
                messages.success(request, "Ответы обновлены.")

        return redirect('questions_list')


class GenerateTestView(View):
    def get(self, request, *args, **kwargs):
        psychology_subject = Subject.objects.get(name="Psychology")

        # Получаем ID вопросов, которые уже использовались в других тестах
        used_question_ids = Test.objects.values_list('questions', flat=True)

        # Получаем все вопросы по предмету, исключая уже использованные
        all_questions = Question.objects.filter(subject=psychology_subject).exclude(id__in=used_question_ids)

        questions_count = min(30, all_questions.count())
        selected_questions = sample(list(all_questions), questions_count)

        test_name = "Психология Test {}".format(Test.objects.count() + 1)
        test = Test.objects.create(name=test_name)
        test.questions.add(*selected_questions)

        return redirect('home')


class GeneratePhilosophyTestView(View):
    def get(self, request, *args, **kwargs):
        # Получение объекта предмета "Философия"
        philosophy_subject = Subject.objects.get(name="Philosophy")

        questions = Question.objects.filter(subject=philosophy_subject).order_by('?')[:30]

        test = Test.objects.create(name="Философия Test {}".format(Test.objects.count() + 1))
        test.questions.add(*questions)

        return redirect('home')


class TestDetailView(DetailView):
    model = Test
    template_name = 'test_detail.html'
    context_object_name = 'test'


class SubmitTestView(View):
    def post(self, request, *args, **kwargs):
        user = request.user
        for key, value in request.POST.items():
            if key.startswith('question_'):
                question_id = key.split('_')[1]
                selected_answer_id = value
                UserAnswer.objects.create(
                    user=user,
                    question_id=question_id,
                    selected_answer_id=selected_answer_id
                )
        test_id = self.kwargs.get('test_id')
        return redirect('test_result', test_id=test_id)


class TestResultView(TemplateView):
    template_name = 'test_result.html'

    def get_context_data(self, **kwargs):
        context = super(TestResultView, self).get_context_data(**kwargs)
        test_id = self.kwargs['test_id']
        user_answers = UserAnswer.objects.filter(user=self.request.user, question__test=test_id)

        total_questions = Test.objects.get(id=test_id).questions.count()
        correct_answers_count = 0

        questions_with_answers = []
        for user_answer in user_answers:
            print(user_answer)
            correct_answer = CorrectAnswer.objects.get(question=user_answer.question)
            is_correct = user_answer.selected_answer.text == correct_answer.answer.text
            if is_correct:
                correct_answers_count += 1

            questions_with_answers.append({
                'question': user_answer.question,
                'selected_answer': user_answer.selected_answer,
                'correct_answer': correct_answer.answer,
                'is_correct': is_correct,
            })

        context['total_questions'] = total_questions
        context['correct_answers_count'] = correct_answers_count
        context['questions_with_answers'] = questions_with_answers
        context['test_id'] = test_id
        return context


def generate_docx(request, test_id):
    test = Test.objects.get(id=test_id)
    document = Document()
    document.add_heading(test.name, 0)

    user_answers = UserAnswer.objects.filter(user=request.user, question__test=test_id)
    for user_answer in user_answers:
        correct_answer = CorrectAnswer.objects.get(question=user_answer.question)
        if user_answer.selected_answer == correct_answer.answer:
            document.add_paragraph(user_answer.question.text, style='ListNumber')

            # Создание абзаца для ответа с жирным шрифтом
            p = document.add_paragraph(style='ListBullet')
            p.add_run(user_answer.selected_answer.text).bold = True

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=correct_answers.docx'
    document.save(response)

    return response


def generate_docx_with_correct_answers_by_subject(request, subject_name='Psychology'):
    document = Document()
    document.add_heading(f'Список Вопросов и Правильных Ответов по Предмету: {subject_name}', 0)

    subject = Subject.objects.get(name=subject_name)

    questions = Question.objects.filter(subject=subject)
    for question in questions:
        correct_answer = CorrectAnswer.objects.get(question=question)

        document.add_paragraph(question.text, style='ListNumber')

        p = document.add_paragraph(style='ListBullet')
        p.add_run(correct_answer.answer.text).bold = True

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{subject_name}_correct_answers.docx"'
    document.save(response)

    return response


def generate_docx_with_questions_and_answers(request, subject_name='Psychology'):
    document = Document()
    document.add_heading(f'Список Вопросов и Ответов по Предмету: {subject_name}', 0)

    # Fetch the subject by name
    subject = Subject.objects.get(name=subject_name)

    # Fetch questions related to the subject
    questions = Question.objects.filter(subject=subject)
    for question in questions:
        # Add the question text
        document.add_paragraph(question.text, style='ListNumber')

        # Retrieve and add all answer options for the question
        answer_options = AnswerOption.objects.filter(question=question)
        for answer_option in answer_options:
            p = document.add_paragraph(style='ListBullet')
            p.add_run(answer_option.text)

    # Configure the response to return a Word document
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{subject_name}_questions_and_answers.docx"'
    document.save(response)

    return response